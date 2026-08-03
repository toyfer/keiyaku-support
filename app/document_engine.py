"""帳票エンジン: 案件 → ひな形入力シートQ列。変更1–5・物品マップ対応。"""
from __future__ import annotations

import shutil
from pathlib import Path

from .config import DATA_DIR, OUTPUT_DIR, TEMPLATES_DIR
from .fields import (
    CHANGE_MAX, TEMPLATE_FILES, all_change_write_cells, cell_map_for,
    change_map, formula_cells_for, map_type_for,
)


def _placeholder_re():
    import re
    return re.compile(r"\{\{\s*(\w+)\s*\}\}")


def resolve_template(case_type: str) -> Path | None:
    rel = TEMPLATE_FILES.get(case_type) or TEMPLATE_FILES.get("業務" if case_type != "工事" else "工事")
    candidates = [DATA_DIR / rel, TEMPLATES_DIR / Path(rel).name, DATA_DIR / "templates" / Path(rel).name]
    name_map = {
        "工事": ["hinagata_koji.xlsx"], "業務": ["hinagata_gyomu.xlsx"],
        "建設業務": ["hinagata_kensetsu_gyomu.xlsx"], "物品": ["hinagata_buppin.xlsx"],
    }
    for n in name_map.get(case_type, []):
        candidates += [TEMPLATES_DIR / n, DATA_DIR / "templates" / n]
    for p in candidates:
        if p and p.exists():
            return p
    return None


def _set_cell(ws, addr: str, value, skip_formula=True, allow_clear=False, protect=None):
    from openpyxl.cell.cell import MergedCell
    from openpyxl.utils import get_column_letter
    protect = protect or set()
    cell = ws[addr]
    if isinstance(cell, MergedCell):
        for mr in ws.merged_cells.ranges:
            try:
                if cell.coordinate in mr or addr in mr:
                    addr = f"{get_column_letter(mr.min_col)}{mr.min_row}"
                    cell = ws[addr]
                    break
            except Exception:
                if addr in mr:
                    addr = f"{get_column_letter(mr.min_col)}{mr.min_row}"
                    cell = ws[addr]
                    break
    if value is None:
        return False
    if value == "" and not allow_clear:
        return False
    if skip_formula and isinstance(cell.value, str) and cell.value.startswith("="):
        if addr in protect:
            return False
        if "VLOOKUP" in cell.value or "IF(" in cell.value:
            cell.value = value
            return True
        return False
    try:
        cell.value = None if value == "" else value
    except AttributeError:
        return False
    return True


def _write_map(ws, cmap, values, protect, written, skipped):
    for key, cell in cmap.items():
        if cell in protect:
            skipped.append((key, cell, "formula-protected"))
            continue
        if key not in values:
            continue
        val = values.get(key)
        if val is None:
            val = ""
        ok = _set_cell(ws, cell, val if val != "" else "", True, True, protect)
        if ok:
            written.append((key, cell, val))
        else:
            skipped.append((key, cell, "skip"))


def _clear_cells(ws, cells, protect, written):
    for cell in cells:
        if cell in protect:
            continue
        if _set_cell(ws, cell, "", True, True, protect):
            written.append(("__clear__", cell, ""))


def fill_excel_template(src, dst, values, case_type="工事", changes=None) -> Path:
    from openpyxl import load_workbook
    src, dst = Path(src), Path(dst)
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    wb = load_workbook(dst)
    mt = map_type_for(case_type)
    # 変更マップ用の種別（物品は業務ストライド）
    ch_type = "工事" if mt == "工事" else "業務"
    cmap = cell_map_for(case_type)
    protect = formula_cells_for(case_type)
    sheet_name = "入力シート" if "入力シート" in wb.sheetnames else wb.sheetnames[0]
    ws = wb[sheet_name]
    written, skipped = [], []
    used_nos = set()
    change_rows = changes or values.get("_changes") or []
    if change_rows:
        _clear_cells(ws, all_change_write_cells(ch_type), protect, written)
        for ch in change_rows:
            try:
                n = int(ch.get("change_no") or 0)
            except (TypeError, ValueError):
                continue
            if not (1 <= n <= CHANGE_MAX) or ch.get("status") == "cancelled":
                continue
            used_nos.add(n)
            cm = change_map(ch_type, n)
            ch_vals = {
                "change_notice": ch.get("change_notice"),
                "change_explain": ch.get("change_explain"),
                "change_kessai_date": ch.get("change_kessai_date"),
                "change_date": ch.get("change_date"),
                "change_dir": ch.get("change_dir"),
                "change_amount_ex": ch.get("change_amount_ex") if ch.get("change_amount_ex") is not None else ch.get("amount_delta_ex"),
                "change_end_date": ch.get("change_end_date") or ch.get("new_end_date"),
                "budget_allocated": ch.get("budget_allocated"),
                "budget_executed": ch.get("budget_executed"),
                "design_delta_ex": ch.get("design_delta_ex"),
            }
            _write_map(ws, cm, ch_vals, protect, written, skipped)
    base_map = dict(cmap)
    if used_nos:
        for k in ("change_notice", "change_explain", "change_kessai_date", "change_date",
                  "change_dir", "change_amount_ex", "change_end_date"):
            base_map.pop(k, None)
    _write_map(ws, base_map, values, protect, written, skipped)
    pattern = _placeholder_re()
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and "{{" in cell.value:
                    cell.value = pattern.sub(lambda m: str(values.get(m.group(1), "") or ""), cell.value)
    try:
        wb.calculation.fullCalcOnLoad = True
    except Exception:
        pass
    wb.save(dst)
    wb.close()
    values["_write_log"] = written
    values["_skip_log"] = skipped
    values["_change_nos"] = sorted(used_nos)
    return dst


def generate_from_case(case, values, out_dir=None, changes=None) -> Path:
    ctype = case.get("case_type") or "工事"
    tpl = resolve_template(ctype) or resolve_template("工事" if ctype == "工事" else "業務")
    if not tpl:
        raise FileNotFoundError(
            "ひな形Excelが見つかりません。data/templates/ に hinagata_*.xlsx を配置してください（分割パッケージB）。"
        )
    out_dir = Path(out_dir or OUTPUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_no = str(case.get("case_no") or "case").replace("/", "-")
    title = str(case.get("title") or "案件")[:40]
    for ch in '\\/:*?"<>|':
        title = title.replace(ch, "_")
    out_name = f"{safe_no}_{title}_{ctype}ひな形.xlsx"
    chs = changes if changes is not None else values.get("_changes")
    return fill_excel_template(tpl, out_dir / out_name, values, case_type=ctype, changes=chs)


def generate_form(template, values, out_dir=None, output_name=None, case_type="工事") -> Path:
    template = Path(template)
    out_dir = Path(out_dir or OUTPUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    dst = out_dir / (output_name or template.name)
    suffix = template.suffix.lower()
    if suffix in (".xlsx", ".xlsm"):
        return fill_excel_template(template, dst, values, case_type=case_type)
    if suffix == ".docx":
        return fill_word_docx(template, dst, values)
    if suffix == ".doc":
        raise RuntimeError("旧形式.docは未対応です。Excelひな形を使ってください。")
    raise RuntimeError(f"未対応の拡張子です: {suffix}")


def fill_word_docx(src, dst, values) -> Path:
    from docx import Document
    src, dst = Path(src), Path(dst)
    dst.parent.mkdir(parents=True, exist_ok=True)
    pattern = _placeholder_re()
    def replace_para(p):
        for run in p.runs:
            if "{{" in (run.text or ""):
                run.text = pattern.sub(lambda m: str(values.get(m.group(1), "") or ""), run.text)
    doc = Document(src)
    for para in doc.paragraphs:
        replace_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_para(para)
    doc.save(dst)
    return dst


def inspect_template(path) -> dict:
    from openpyxl import load_workbook
    import re
    from collections import Counter
    path = Path(path)
    wb = load_workbook(path, data_only=False)
    info = {"file": path.name, "sheets": wb.sheetnames, "has_input": "入力シート" in wb.sheetnames,
            "input_labels": [], "ref_columns": Counter(), "broken_names": []}
    try:
        for dn in wb.defined_names:
            txt = getattr(wb.defined_names[dn], "attr_text", "") or ""
            if "#REF!" in txt or "#N/A" in txt:
                info["broken_names"].append(f"{dn}={txt}")
    except Exception:
        pass
    if info["has_input"]:
        ws = wb["入力シート"]
        for r in range(1, min(140, (ws.max_row or 140) + 1)):
            lab, q = ws.cell(r, 9).value, ws.cell(r, 17).value
            if lab or q is not None:
                qshow = f"[数式]{q[:40]}" if isinstance(q, str) and q.startswith("=") else q
                info["input_labels"].append({"row": r, "label": lab, "Q": qshow})
    pat = re.compile(r"入力シート[!！]?\$?([A-Z]{1,3})\$?(\d+)", re.I)
    for name in wb.sheetnames:
        if name in ("入力シート", "業者登録一覧"):
            continue
        s = wb[name]
        for row in s.iter_rows(max_row=min(s.max_row or 80, 100), max_col=min(s.max_column or 40, 40)):
            for cell in row:
                v = cell.value
                if isinstance(v, str) and "入力シート" in v:
                    for m in pat.finditer(v):
                        info["ref_columns"][m.group(1)] += 1
    info["ref_columns"] = dict(info["ref_columns"].most_common(10))
    wb.close()
    return info
